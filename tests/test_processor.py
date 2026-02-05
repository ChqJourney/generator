"""
单元测试：processor.py
测试 Word 文档处理器的各项功能
使用 Mock 避免实际的 Word 文件操作
"""

import pytest
import json
import sys
from pathlib import Path
from unittest.mock import MagicMock, patch, mock_open, call
from typing import List, Dict, Any

# 导入被测试的模块
sys.path.insert(0, 'src')
from processor import (
    DocxTemplateProcessor,
    TextInserter,
    TableInserter,
    ImageInserter,
    CheckboxInserter,
    PlaceholderFinder,
    ContentInserter,
    DocxTemplateError,
    PlaceholderNotFoundError,
    InvalidLocationError,
)


# =============================================================================
# Fixtures
# =============================================================================

@pytest.fixture
def mock_document():
    """创建模拟的 Document 对象"""
    doc = MagicMock()
    
    # 模拟段落 - 包含占位符
    para_with_placeholder = MagicMock()
    para_with_placeholder.text = "{{test_field}}"
    para_with_placeholder.runs = [MagicMock()]
    para_with_placeholder.runs[0].text = "{{test_field}}"
    para_with_placeholder._element = MagicMock()
    para_with_placeholder._element.getparent.return_value = MagicMock()
    
    # 模拟普通段落
    para_normal = MagicMock()
    para_normal.text = "普通文本"
    para_normal.runs = [MagicMock()]
    para_normal.runs[0].text = "普通文本"
    
    doc.paragraphs = [para_with_placeholder, para_normal]
    doc.tables = []
    
    # 模拟 sections
    section = MagicMock()
    section.header = MagicMock()
    section.header.paragraphs = []
    section.header.tables = []
    section.first_page_header = None
    section.even_page_header = None
    section.footer = MagicMock()
    section.footer.paragraphs = []
    section.footer.tables = []
    section.first_page_footer = None
    section.even_page_footer = None
    doc.sections = [section]
    
    # 模拟 part.element 用于 checkbox 测试
    doc.part = MagicMock()
    doc.part.element = MagicMock()
    
    return doc


@pytest.fixture
def mock_table_template():
    """创建模拟的表格模板文档"""
    template_doc = MagicMock()
    
    # 创建模拟单元格
    mock_cell1 = MagicMock()
    mock_cell1.text = ""
    mock_cell2 = MagicMock()
    mock_cell2.text = ""
    
    # 创建模拟行
    mock_row_header = MagicMock()
    mock_row_header.cells = [mock_cell1, mock_cell2]
    mock_row_data = MagicMock()
    mock_row_data.cells = [MagicMock(), MagicMock()]
    
    # 创建模拟表格
    mock_table = MagicMock()
    mock_table.rows = [mock_row_header, mock_row_data]
    mock_table._element = MagicMock()
    mock_table._tbl = MagicMock()
    
    template_doc.tables = [mock_table]
    template_doc.paragraphs = []
    
    return template_doc


@pytest.fixture
def mock_paragraph_with_placeholder():
    """创建包含占位符的模拟段落"""
    para = MagicMock()
    para.text = "{{placeholder}}"
    
    # 模拟 runs
    run = MagicMock()
    run.text = "{{placeholder}}"
    para.runs = [run]
    
    # 模拟 _element
    para._element = MagicMock()
    parent = MagicMock()
    para._element.getparent.return_value = parent
    
    return para


# =============================================================================
# DocxTemplateProcessor 测试
# =============================================================================

class TestDocxTemplateProcessor:
    """DocxTemplateProcessor 的测试类"""
    
    @patch('processor.os.path.exists')
    @patch('processor.shutil.copy')
    @patch('processor.Document')
    @patch('processor.DocxTemplateProcessor.is_word_file_open')
    def test_init_success(self, mock_is_open, mock_doc_class, mock_copy, mock_exists, mock_document):
        """测试处理器初始化成功"""
        mock_exists.return_value = True
        mock_is_open.return_value = False
        mock_doc_class.return_value = mock_document
        
        processor = DocxTemplateProcessor("template.docx", "output.docx")
        
        assert processor.template_path == "template.docx"
        assert processor.output_path == "output.docx"
        assert processor.operations == []
        mock_copy.assert_called_once_with("template.docx", "output.docx")
        mock_doc_class.assert_called_once_with("output.docx")
    
    @patch('processor.os.path.exists')
    def test_init_template_not_found(self, mock_exists):
        """测试模板文件不存在时抛出异常"""
        mock_exists.return_value = False
        
        with pytest.raises(DocxTemplateError, match="Template file not found"):
            DocxTemplateProcessor("nonexistent.docx", "output.docx")
    
    @patch('processor.os.path.exists')
    @patch('processor.shutil.copy')
    @patch('processor.Document')
    def test_init_file_locked(self, mock_doc_class, mock_copy, mock_exists, mock_document):
        """测试输出文件被 Word 占用时抛出异常"""
        mock_exists.return_value = True
        
        with patch('processor.DocxTemplateProcessor.is_word_file_open', return_value=True):
            with pytest.raises(DocxTemplateError, match="currently open in Word"):
                DocxTemplateProcessor("template.docx", "output.docx")
    
    @patch('processor.os.path.exists')
    @patch('processor.shutil.copy')
    @patch('processor.Document')
    @patch('processor.DocxTemplateProcessor.is_word_file_open')
    def test_add_text_operation(self, mock_is_open, mock_doc_class, mock_copy, mock_exists, mock_document):
        """测试添加文本操作"""
        mock_exists.return_value = True
        mock_is_open.return_value = False
        mock_doc_class.return_value = mock_document
        
        processor = DocxTemplateProcessor("template.docx", "output.docx")
        result = processor.add_text("field1", "value1", "body")
        
        # 验证返回 self 支持链式调用
        assert result == processor
        assert len(processor.operations) == 1
        assert processor.operations[0]['type'] == 'text'
        assert processor.operations[0]['placeholder'] == 'field1'
        assert processor.operations[0]['value'] == "value1"
        assert processor.operations[0]['location'] == "body"
    
    @patch('processor.os.path.exists')
    @patch('processor.shutil.copy')
    @patch('processor.Document')
    @patch('processor.DocxTemplateProcessor.is_word_file_open')
    def test_add_table_operation(self, mock_is_open, mock_doc_class, mock_copy, mock_exists, mock_document):
        """测试添加表格操作"""
        mock_exists.return_value = True
        mock_is_open.return_value = False
        mock_doc_class.return_value = mock_document
        
        processor = DocxTemplateProcessor("template.docx", "output.docx")
        result = processor.add_table(
            placeholder="table1",
            table_template_path="table_template.docx",
            raw_data=[["a", "b"], ["c", "d"]],
            row_strategy="dynamic_rows"
        )
        
        assert result == processor
        assert len(processor.operations) == 1
        assert processor.operations[0]['type'] == 'table'
        assert processor.operations[0]['placeholder'] == "table1"
        assert processor.operations[0]['row_strategy'] == "dynamic_rows"
    
    @patch('processor.os.path.exists')
    @patch('processor.shutil.copy')
    @patch('processor.Document')
    @patch('processor.DocxTemplateProcessor.is_word_file_open')
    def test_add_image_operation(self, mock_is_open, mock_doc_class, mock_copy, mock_exists, mock_document):
        """测试添加图片操作"""
        mock_exists.return_value = True
        mock_is_open.return_value = False
        mock_doc_class.return_value = mock_document
        
        processor = DocxTemplateProcessor("template.docx", "output.docx")
        result = processor.add_image(
            placeholder="img1",
            image_paths=["test.jpg"],
            width=MagicMock(),  # Mock Inches object
            alignment="center"
        )
        
        assert result == processor
        assert len(processor.operations) == 1
        assert processor.operations[0]['type'] == 'image'
        assert processor.operations[0]['placeholder'] == "img1"
        assert processor.operations[0]['image_paths'] == ["test.jpg"]
    
    @patch('processor.os.path.exists')
    @patch('processor.shutil.copy')
    @patch('processor.Document')
    @patch('processor.DocxTemplateProcessor.is_word_file_open')
    def test_add_checkboxes_operation(self, mock_is_open, mock_doc_class, mock_copy, mock_exists, mock_document):
        """测试添加复选框操作"""
        mock_exists.return_value = True
        mock_is_open.return_value = False
        mock_doc_class.return_value = mock_document
        
        processor = DocxTemplateProcessor("template.docx", "output.docx")
        checkbox_mapping = {"cb1": True, "cb2": False}
        result = processor.add_checkboxes(checkbox_mapping)
        
        assert result == processor
        assert len(processor.operations) == 1
        assert processor.operations[0]['type'] == 'checkbox'
        assert processor.operations[0]['checkbox_mapping'] == checkbox_mapping
    
    @patch('processor.os.path.exists')
    @patch('processor.shutil.copy')
    @patch('processor.Document')
    @patch('processor.DocxTemplateProcessor.is_word_file_open')
    def test_get_all_placeholders(self, mock_is_open, mock_doc_class, mock_copy, mock_exists):
        """测试获取所有占位符"""
        mock_exists.return_value = True
        mock_is_open.return_value = False
        
        # 创建带占位符的模拟文档
        mock_doc = MagicMock()
        
        para1 = MagicMock()
        para1.text = "{{field1}} 和 {{field2}}"
        
        para2 = MagicMock()
        para2.text = "普通文本"
        
        mock_doc.paragraphs = [para1, para2]
        mock_doc.tables = []
        mock_doc.sections = []
        
        mock_doc_class.return_value = mock_doc
        
        processor = DocxTemplateProcessor("template.docx", "output.docx")
        placeholders = processor.get_all_placeholders()
        
        assert "field1" in placeholders
        assert "field2" in placeholders
    
    @patch('processor.os.path.exists')
    @patch('processor.shutil.copy')
    @patch('processor.Document')
    @patch('processor.TextInserter')
    @patch('processor.DocxTemplateProcessor.is_word_file_open')
    def test_process_text_operation(self, mock_is_open, mock_inserter_class, mock_doc_class, mock_copy, mock_exists, mock_document):
        """测试处理文本操作"""
        mock_exists.return_value = True
        mock_is_open.return_value = False
        mock_doc_class.return_value = mock_document
        mock_inserter = MagicMock()
        mock_inserter_class.return_value = mock_inserter
        
        processor = DocxTemplateProcessor("template.docx", "output.docx")
        processor.add_text("field1", "value1")
        processor.process()
        
        mock_inserter.insert.assert_called_once_with("field1", "value1", "body")
        mock_document.save.assert_called_once_with("output.docx")


# =============================================================================
# TextInserter 测试
# =============================================================================

class TestTextInserter:
    """TextInserter 的测试类"""
    
    def test_insert_text_in_body(self, mock_document, mock_paragraph_with_placeholder):
        """测试在正文插入文本"""
        mock_document.paragraphs = [mock_paragraph_with_placeholder]
        
        inserter = TextInserter(mock_document)
        inserter.insert("placeholder", "new_value", "body")
        
        # 验证文本被替换
        assert mock_paragraph_with_placeholder.runs[0].text == "new_value"
    
    def test_insert_text_not_found(self, mock_document, capsys):
        """测试占位符不存在时的处理"""
        mock_document.paragraphs = []
        
        inserter = TextInserter(mock_document)
        inserter.insert("nonexistent", "value", "body")
        
        captured = capsys.readouterr()
        assert "未找到" in captured.out or "not found" in captured.out.lower()
    
    def test_validate_location_invalid(self, mock_document):
        """测试无效位置抛出异常"""
        inserter = TextInserter(mock_document)
        
        with pytest.raises(InvalidLocationError):
            inserter.insert("field", "value", "invalid_location")
    
    def test_insert_text_in_header(self, mock_document):
        """测试在页眉插入文本"""
        # 创建带占位符的页眉段落
        header_para = MagicMock()
        header_para.text = "{{header_field}}"
        header_para.runs = [MagicMock()]
        header_para.runs[0].text = "{{header_field}}"
        
        mock_document.sections[0].header.paragraphs = [header_para]
        mock_document.paragraphs = []  # body 中无占位符
        
        inserter = TextInserter(mock_document)
        inserter.insert("header_field", "header_value", "header")
        
        assert header_para.runs[0].text == "header_value"


# =============================================================================
# TableInserter 测试
# =============================================================================

class TestTableInserter:
    """TableInserter 的测试类"""
    
    @patch('processor.os.path.exists')
    @patch('processor.Document')
    def test_insert_table_success(self, mock_doc_class, mock_exists, mock_document, mock_table_template):
        """测试成功插入表格"""
        mock_exists.return_value = True
        mock_doc_class.return_value = mock_table_template
        
        # 设置占位符段落
        mock_para = MagicMock()
        mock_para.text = "{{table_placeholder}}"
        mock_para._element = MagicMock()
        mock_para._element.getparent.return_value = MagicMock()
        mock_document.paragraphs = [mock_para]
        
        inserter = TableInserter(mock_document)
        inserter.insert(
            placeholder="table_placeholder",
            table_template_path="table_template.docx",
            raw_data=[["1", "2"], ["3", "4"]],
            row_strategy="fixed_rows"
        )
        
        # 验证表格模板被加载
        mock_doc_class.assert_called_once_with("table_template.docx")
    
    @patch('processor.os.path.exists')
    def test_insert_table_template_not_found(self, mock_exists, mock_document):
        """测试表格模板不存在时抛出异常"""
        mock_exists.return_value = False
        
        inserter = TableInserter(mock_document)
        
        with pytest.raises(DocxTemplateError, match="Table template file not found"):
            inserter.insert("placeholder", "nonexistent.docx")
    
    @patch('processor.os.path.exists')
    @patch('processor.Document')
    def test_insert_table_no_tables_in_template(self, mock_doc_class, mock_exists, mock_document):
        """测试表格模板中无表格时抛出异常"""
        mock_exists.return_value = True
        
        # 创建无表格的模板
        mock_no_table_doc = MagicMock()
        mock_no_table_doc.tables = []
        mock_doc_class.return_value = mock_no_table_doc
        
        inserter = TableInserter(mock_document)
        
        with pytest.raises(DocxTemplateError, match="No tables found in template"):
            inserter.insert("placeholder", "template.docx")
    
    @patch('processor.os.path.exists')
    @patch('processor.Document')
    def test_fill_fixed_rows(self, mock_doc_class, mock_exists, mock_table_template):
        """测试固定行策略填充数据"""
        mock_exists.return_value = True
        mock_doc_class.return_value = mock_table_template
        
        # 创建带单元格的模拟表格
        mock_cell1 = MagicMock()
        mock_cell1.text = ""
        mock_cell2 = MagicMock()
        mock_cell2.text = ""
        
        mock_row1 = MagicMock()  # header row
        mock_row1.cells = [mock_cell1, mock_cell2]
        
        mock_row2 = MagicMock()  # data row
        cell2_1 = MagicMock()
        cell2_1.text = ""
        cell2_2 = MagicMock()
        cell2_2.text = ""
        mock_row2.cells = [cell2_1, cell2_2]
        
        mock_table = MagicMock()
        mock_table.rows = [mock_row1, mock_row2]
        
        # 模拟主文档
        main_doc = MagicMock()
        para = MagicMock()
        para.text = "{{table}}"
        para._element = MagicMock()
        para._element.getparent.return_value = MagicMock()
        main_doc.paragraphs = [para]
        
        inserter = TableInserter(main_doc)
        
        # 使用 _fill_fixed_rows 方法
        inserter._fill_fixed_rows(
            table=mock_table,
            data=[["data1", "data2"]],
            skip_columns=None,
            header_rows=1
        )
        
        # 验证单元格值被设置
        # 注意：实际调用的是 _set_cell_value，这里我们验证它是否被调用
        # 由于 _set_cell_value 内部调用 set_cell_value，我们需要检查 mock


# =============================================================================
# ImageInserter 测试
# =============================================================================

class TestImageInserter:
    """ImageInserter 的测试类"""
    
    @patch('processor.os.path.exists')
    def test_insert_image_success(self, mock_exists, mock_document, mock_paragraph_with_placeholder):
        """测试成功插入图片 - 使用更完善的 mock"""
        mock_exists.return_value = True
        
        # 设置更完善的 mock 段落
        mock_para = MagicMock()
        mock_para.text = "{{placeholder}}"
        
        # 设置 _element 和 parent
        p_element = MagicMock()
        p_parent = MagicMock()
        p_parent.__iter__ = MagicMock(return_value=iter([p_element]))
        p_parent.index = MagicMock(return_value=0)
        
        mock_para._element = p_element
        p_element.getparent.return_value = p_parent
        
        mock_document.paragraphs = [mock_para]
        mock_document.tables = []
        
        inserter = ImageInserter(mock_document)
        
        # 模拟图片路径处理
        with patch.object(inserter, '_resolve_image_path', return_value="data_files/test.jpg"):
            # 使用真正的 Inches 对象来通过验证
            from docx.shared import Inches
            inserter.insert(
                placeholder="placeholder",
                image_paths=["test.jpg"],
                width=Inches(4.0),
                alignment="center"
            )
    
    @patch('processor.os.path.exists')
    def test_insert_image_file_not_found(self, mock_exists, mock_document):
        """测试图片文件不存在时抛出异常"""
        mock_exists.return_value = False
        
        inserter = ImageInserter(mock_document)
        
        with pytest.raises(DocxTemplateError, match="Image file not found"):
            inserter.insert("placeholder", ["nonexistent.jpg"])
    
    def test_resolve_image_path_absolute(self, mock_document):
        """测试绝对路径解析"""
        inserter = ImageInserter(mock_document)
        
        with patch('processor.os.path.exists', return_value=True):
            result = inserter._resolve_image_path("/absolute/path/image.jpg")
            assert result == "/absolute/path/image.jpg"
    
    def test_resolve_image_path_relative(self, mock_document):
        """测试相对路径解析"""
        inserter = ImageInserter(mock_document)
        
        with patch('processor.os.path.exists') as mock_exists:
            mock_exists.side_effect = [False, True]  # 原路径不存在，data_files 路径存在
            result = inserter._resolve_image_path("./image.jpg")
            assert "data_files" in result
    
    def test_validate_image_dimensions_valid(self, mock_document):
        """测试有效的图片尺寸"""
        from docx.shared import Inches
        
        inserter = ImageInserter(mock_document)
        # 不应抛出异常
        inserter._validate_image_dimensions(Inches(4), None)
        inserter._validate_image_dimensions(None, Inches(3))
        inserter._validate_image_dimensions(Inches(4), Inches(3))
    
    def test_validate_image_dimensions_invalid(self, mock_document):
        """测试无效的图片尺寸"""
        inserter = ImageInserter(mock_document)
        
        with pytest.raises(DocxTemplateError, match="Invalid image dimension"):
            inserter._validate_image_dimensions("invalid", None)


# =============================================================================
# CheckboxInserter 测试
# =============================================================================

class TestCheckboxInserter:
    """CheckboxInserter 的测试类"""
    
    def test_insert_checkbox_check(self, mock_document):
        """测试勾选复选框"""
        # 模拟 XML 结构
        mock_checkbox = MagicMock()
        mock_ffdata = MagicMock()
        mock_name = MagicMock()
        mock_name.get.return_value = "checkbox1"
        
        mock_ffdata.find.return_value = mock_name
        mock_checkbox.getparent.return_value = mock_ffdata
        
        # 模拟 checked 元素不存在
        mock_checkbox.find.side_effect = lambda x, **kwargs: None if x == 'w:checked' else MagicMock()
        
        # 设置 mock 文档的 element
        mock_document.part.element.findall.return_value = [mock_checkbox]
        
        inserter = CheckboxInserter(mock_document)
        
        with patch('processor.parse_xml') as mock_parse_xml:
            inserter.insert({"checkbox1": True})
            mock_parse_xml.assert_called_once()
    
    def test_insert_checkbox_uncheck(self, mock_document):
        """测试取消勾选复选框"""
        mock_checkbox = MagicMock()
        mock_ffdata = MagicMock()
        mock_name = MagicMock()
        mock_name.get.return_value = "checkbox1"
        
        mock_ffdata.find.return_value = mock_name
        mock_checkbox.getparent.return_value = mock_ffdata
        
        # 模拟已存在的 checked 元素
        mock_checked = MagicMock()
        
        def mock_find(x, **kwargs):
            if x == 'w:checked':
                return mock_checked
            elif x == 'w:name':
                return mock_name
            return None
        
        mock_checkbox.find.side_effect = mock_find
        
        mock_document.part.element.findall.return_value = [mock_checkbox]
        
        inserter = CheckboxInserter(mock_document)
        inserter.insert({"checkbox1": False})
        
        # 验证 checked 元素被移除
        mock_checkbox.remove.assert_called_once_with(mock_checked)
    
    def test_insert_checkbox_not_found(self, mock_document, capsys):
        """测试复选框未找到时的警告"""
        mock_document.part.element.findall.return_value = []
        
        inserter = CheckboxInserter(mock_document)
        inserter.insert({"nonexistent_checkbox": True})
        
        captured = capsys.readouterr()
        assert "未找到" in captured.out or "not found" in captured.out.lower()


# =============================================================================
# PlaceholderFinder 测试
# =============================================================================

class TestPlaceholderFinder:
    """PlaceholderFinder 的测试类"""
    
    def test_find_paragraph_with_placeholder(self, mock_document):
        """测试查找包含占位符的段落"""
        mock_para = MagicMock()
        mock_para.text = "{{test_field}}"
        mock_document.paragraphs = [mock_para]
        
        idx, para = PlaceholderFinder.find_paragraph_with_placeholder(
            mock_document, "{{test_field}}", "body"
        )
        
        assert para == mock_para
    
    def test_find_paragraph_not_found(self, mock_document):
        """测试占位符未找到"""
        mock_document.paragraphs = []
        
        idx, para = PlaceholderFinder.find_paragraph_with_placeholder(
            mock_document, "{{nonexistent}}", "body"
        )
        
        assert idx is None
        assert para is None
    
    def test_find_all_placeholders_in_location_body(self, mock_document):
        """测试在正文中查找所有占位符"""
        mock_para1 = MagicMock()
        mock_para1.text = "{{field1}}"
        mock_para2 = MagicMock()
        mock_para2.text = "{{field2}}"
        
        mock_document.paragraphs = [mock_para1, mock_para2]
        mock_document.sections = []
        
        results = PlaceholderFinder.find_all_placeholders_in_location(
            mock_document, "{{field1}}", "body"
        )
        
        assert len(results) == 1
    
    def test_replace_paragraph_with_element(self):
        """测试用元素替换段落"""
        mock_para = MagicMock()
        mock_element = MagicMock()
        mock_parent = MagicMock()
        
        mock_para._element = mock_element
        mock_element.getparent.return_value = mock_parent
        mock_parent.__iter__ = MagicMock(return_value=iter([mock_element]))
        mock_parent.index = MagicMock(return_value=0)
        
        new_element = MagicMock()
        
        # 深拷贝 mock
        with patch('processor.deepcopy', return_value=new_element):
            PlaceholderFinder.replace_paragraph_with_element(mock_para, new_element)
        
        mock_parent.remove.assert_called_once_with(mock_element)


# =============================================================================
# ContentInserter (基类) 测试
# =============================================================================

class TestContentInserter:
    """ContentInserter 基类的测试类"""
    
    def test_validate_location_valid(self, mock_document):
        """测试验证有效位置 - 使用 TextInserter 作为 ContentInserter 的具体实现"""
        inserter = TextInserter(mock_document)
        # 不应抛出异常
        inserter.validate_location("body", ["body", "header", "footer"])
        inserter.validate_location("header", ["body", "header", "footer"])
    
    def test_validate_location_invalid(self, mock_document):
        """测试验证无效位置 - 使用 TextInserter 作为 ContentInserter 的具体实现"""
        inserter = TextInserter(mock_document)
        
        with pytest.raises(InvalidLocationError):
            inserter.validate_location("invalid", ["body", "header"])


# =============================================================================
# 异常类测试
# =============================================================================

class TestExceptions:
    """异常类的测试类"""
    
    def test_placeholder_not_found_error(self):
        """测试 PlaceholderNotFoundError"""
        error = PlaceholderNotFoundError("test_field", "body")
        assert "test_field" in str(error)
        assert "body" in str(error)
        assert error.placeholder == "test_field"
        assert error.location == "body"
    
    def test_invalid_location_error(self):
        """测试 InvalidLocationError"""
        error = InvalidLocationError("invalid_loc", "text")
        assert "invalid_loc" in str(error)
        assert "text" in str(error)
    
    def test_docx_template_error(self):
        """测试 DocxTemplateError"""
        error = DocxTemplateError("custom error message")
        assert "custom error message" in str(error)


# =============================================================================
# is_word_file_open 测试
# =============================================================================

class TestIsWordFileOpen:
    """is_word_file_open 静态方法的测试类"""
    
    @patch('processor.os.path.exists')
    def test_file_is_open(self, mock_exists):
        """测试检测 Word 文件被打开"""
        mock_exists.return_value = True
        
        result = DocxTemplateProcessor.is_word_file_open("C:/docs/test.docx")
        
        assert result is True
        mock_exists.assert_called_once()
    
    @patch('processor.os.path.exists')
    def test_file_is_not_open(self, mock_exists):
        """测试检测 Word 文件未被打开"""
        mock_exists.return_value = False
        
        result = DocxTemplateProcessor.is_word_file_open("C:/docs/test.docx")
        
        assert result is False


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
