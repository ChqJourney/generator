#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
调试header/footer内容
详细显示Word文档中所有header/footer的文本内容

用法:
    python tools/debug_header.py your_template.docx
"""

import sys
import re
from docx import Document

def debug_headers_footers(doc_path: str):
    print(f"\n正在分析文档: {doc_path}")
    print("="*60)
    
    doc = Document(doc_path)
    
    # 检查所有可能的placeholder格式
    # 支持字母、数字、下划线、点号（如 {{v.15.1}}）
    patterns = {
        "双括号 {{name}}": re.compile(r'\{\{([\w.]+)\}\}'),
        "双括号带空格 {{ name }}": re.compile(r'\{\{\s*([\w.]+)\s*\}\}'),
        "单括号 {name}": re.compile(r'\{([\w.]+)\}'),
        "方括号 [name]": re.compile(r'\[([\w.]+)\]'),
        "美元符号 $name$": re.compile(r'\$([\w.]+)\$'),
        "百分号 %name%": re.compile(r'%([\w.]+)%'),
    }
    
    for section_idx, section in enumerate(doc.sections, 1):
        print(f"\n{'='*60}")
        print(f"Section {section_idx}")
        print('='*60)
        
        # Header types
        headers = [
            ('Default Header', section.header),
            ('First Page Header', section.first_page_header),
            ('Even Page Header', section.even_page_header),
        ]
        
        for header_name, header in headers:
            if not header:
                continue
                
            print(f"\n--- {header_name} ---")
            
            # 段落
            if header.paragraphs:
                print("段落:")
                for pi, para in enumerate(header.paragraphs, 1):
                    text = para.text.strip()
                    if text:
                        print(f"  [{pi}] {text[:100]}")
                        # 检查各种placeholder格式
                        for pattern_name, pattern in patterns.items():
                            matches = pattern.findall(text)
                            if matches:
                                print(f"      >> 发现{pattern_name}: {matches}")
            
            # 表格
            if header.tables:
                print("表格:")
                for ti, table in enumerate(header.tables, 1):
                    print(f"  表格[{ti}]:")
                    for ri, row in enumerate(table.rows, 1):
                        for ci, cell in enumerate(row.cells, 1):
                            text = cell.text.strip()
                            if text:
                                print(f"    单元格[{ri},{ci}]: {text[:60]}")
                                # 检查各种placeholder格式
                                for pattern_name, pattern in patterns.items():
                                    matches = pattern.findall(text)
                                    if matches:
                                        print(f"        >> 发现{pattern_name}: {matches}")
        
        # Footer types
        footers = [
            ('Default Footer', section.footer),
            ('First Page Footer', section.first_page_footer),
            ('Even Page Footer', section.even_page_footer),
        ]
        
        for footer_name, footer in footers:
            if not footer:
                continue
                
            print(f"\n--- {footer_name} ---")
            
            # 段落
            if footer.paragraphs:
                print("段落:")
                for pi, para in enumerate(footer.paragraphs, 1):
                    text = para.text.strip()
                    if text:
                        print(f"  [{pi}] {text[:100]}")
                        # 检查各种placeholder格式
                        for pattern_name, pattern in patterns.items():
                            matches = pattern.findall(text)
                            if matches:
                                print(f"      >> 发现{pattern_name}: {matches}")
            
            # 表格
            if footer.tables:
                print("表格:")
                for ti, table in enumerate(footer.tables, 1):
                    print(f"  表格[{ti}]:")
                    for ri, row in enumerate(table.rows, 1):
                        for ci, cell in enumerate(row.cells, 1):
                            text = cell.text.strip()
                            if text:
                                print(f"    单元格[{ri},{ci}]: {text[:60]}")
                                # 检查各种placeholder格式
                                for pattern_name, pattern in patterns.items():
                                    matches = pattern.findall(text)
                                    if matches:
                                        print(f"        >> 发现{pattern_name}: {matches}")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法: python tools/debug_header.py your_template.docx")
        sys.exit(1)
    
    debug_headers_footers(sys.argv[1])
