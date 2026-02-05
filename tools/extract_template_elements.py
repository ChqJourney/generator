#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
从Word模板中提取所有placeholder和checkbox
用于生成report_config.json的基础结构

用法:
    python tools/extract_template_elements.py <template.docx> [--output config/extracted_elements.json]
"""

import re
import json
import argparse
from docx import Document
from typing import List, Dict, Set, Tuple
from docx.oxml import parse_xml


def extract_placeholders(doc: Document, pattern_str: str = None) -> Dict[str, List[Dict]]:
    """
    从文档中提取所有placeholder
    
    默认格式: {{placeholder}} 或 {{ placeholder }}
    
    返回按位置分类的placeholder列表:
    {
        "body": [{"name": "xxx", "section_no": 1}],
        "header": [...],
        "footer": [...]
    }
    
    Args:
        doc: Document对象
        pattern_str: 自定义正则表达式（可选）
    """
    # 默认支持 {{name}} 或 {{ name }} 格式
    # 支持字母、数字、下划线、点号、连字符（如 {{v.15.1}}, {{0-180_zone_flux}}）
    if pattern_str:
        pattern = re.compile(pattern_str)
    else:
        pattern = re.compile(r'\{\{\s*([\w.\-]+)\s*\}\}')
    
    placeholders = {"body": [], "header": [], "footer": []}
    # 每个位置单独去重
    body_names = set()
    header_names = set()
    footer_names = set()
    
    def add_placeholder(location: str, name: str, section_no: int = 1):
        """添加placeholder到指定位置，按位置去重"""
        if location == "body" and name not in body_names:
            placeholders["body"].append({"name": name, "section_no": section_no})
            body_names.add(name)
        elif location == "header" and name not in header_names:
            placeholders["header"].append({"name": name, "section_no": section_no})
            header_names.add(name)
        elif location == "footer" and name not in footer_names:
            placeholders["footer"].append({"name": name, "section_no": section_no})
            footer_names.add(name)
    
    # 提取body中的placeholder（属于section 1）
    for para in doc.paragraphs:
        matches = pattern.findall(para.text)
        for match in matches:
            add_placeholder("body", match, section_no=1)
    
    # 提取body表格中的placeholder
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text or ""
                cell_matches = pattern.findall(cell_text)
                
                for match in cell_matches:
                    add_placeholder("body", match, section_no=1)
    
    # 提取header中的placeholder（包括段落和表格）
    for section_idx, section in enumerate(doc.sections, 1):
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header:
                # 从段落中提取
                for para in header.paragraphs:
                    matches = pattern.findall(para.text)
                    for match in matches:
                        add_placeholder("header", match, section_no=section_idx)
                
                # 从表格中提取
                for table in header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            cell_text = cell.text or ""
                            cell_matches = pattern.findall(cell_text)
                            
                            for match in cell_matches:
                                add_placeholder("header", match, section_no=section_idx)
    
    # 提取footer中的placeholder（包括段落和表格）
    for section_idx, section in enumerate(doc.sections, 1):
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer:
                # 从段落中提取
                for para in footer.paragraphs:
                    matches = pattern.findall(para.text)
                    for match in matches:
                        add_placeholder("footer", match, section_no=section_idx)
                
                # 从表格中提取
                for table in footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            cell_text = cell.text or ""
                            cell_matches = pattern.findall(cell_text)
                            
                            for match in cell_matches:
                                add_placeholder("footer", match, section_no=section_idx)
    
    return placeholders


def extract_checkboxes(doc: Document) -> List[Dict]:
    """
    从文档中提取所有checkbox的name属性
    
    返回:
        [{"name": "checkbox_name", "section_no": 1}]
    """
    checkboxes = []
    found_names = set()
    
    root = doc.part.element
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    
    # 查找所有checkbox元素
    checkbox_elements = root.findall('.//w:checkBox', namespaces=ns)
    
    for checkbox in checkbox_elements:
        ffdata = checkbox.getparent()
        if ffdata is not None:
            name_elem = ffdata.find('w:name', namespaces=ns)
            if name_elem is not None:
                w_ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
                field_name = name_elem.get(w_ns + 'val')
                
                if field_name and field_name not in found_names:
                    checkboxes.append({
                        "name": field_name,
                        "section_no": 1  # checkbox默认归于section 1
                    })
                    found_names.add(field_name)
    
    return checkboxes


def smart_infer_field_info(name: str, section_no: int = 1) -> Dict:
    """
    根据placeholder名称智能推断字段类型和数据源
    """
    name_lower = name.lower()
    
    # 推断类型
    field_type = "text"  # 默认类型
    
    # 图像类型关键词
    image_keywords = ['image', 'img', 'photo', 'picture', 'pic', 'logo', 'signature', 'sign']
    if any(kw in name_lower for kw in image_keywords):
        field_type = "image"
    
    # 表格类型关键词
    table_keywords = ['table', 'data', 'list', 'items', 'records', 'photometric', 'test_data']
    if any(kw in name_lower for kw in table_keywords):
        field_type = "table"
    
    # checkbox类型（本身是checkbox）
    if name_lower.startswith('cb_') or name_lower.startswith('chk_'):
        field_type = "checkbox"
    
    # 推断数据源
    source_prefix = "extracted_data"
    if any(kw in name_lower for kw in ['report', 'issue_date', 'applicant', 'product', 'manufacturer', 'sample']):
        source_prefix = "metadata"
    elif any(kw in name_lower for kw in ['energy', 'efficacy', 'class', 'rating', 'calculate']):
        source_prefix = "calculated_data"
    
    # 推断可能的计算函数
    suggested_function = None
    suggested_args = []
    
    if 'energy_class' in name_lower or 'class_rating' in name_lower:
        suggested_function = "calculate_energy_class_rating"
        suggested_args = ["extracted_data.rated_wattage", "extracted_data.useful_luminous_flux"]
    elif 'efficacy' in name_lower:
        suggested_function = "calculate_energy_efficacy"
        suggested_args = ["extracted_data.rated_wattage", "extracted_data.useful_luminous_flux"]
    elif 'percentage' in name_lower or 'percent' in name_lower:
        suggested_function = "calculate_percentage"
        suggested_args = ["extracted_data.value1", "extracted_data.value2"]
    
    result = {
        "template_field": name,
        "source_field": f"{source_prefix}.{name}",
        "type": field_type,
        "inferred_source": source_prefix,
        "section_no": section_no
    }
    
    if suggested_function:
        result["suggested_function"] = suggested_function
        result["suggested_args"] = suggested_args
    
    # 根据类型添加额外配置
    if field_type == "image":
        result["width"] = 4.0
        result["alignment"] = "center"
    elif field_type == "table":
        result["row_strategy"] = "fixed_rows"
        result["header_rows"] = 1
        result["table_template_path"] = "report_templates/table_template.docx"
    
    return result


def generate_field_mappings(placeholders: Dict, checkboxes: List) -> List[Dict]:
    """
    生成field_mappings列表
    合并所有位置的placeholder，并标记它们出现的位置
    """
    mappings = []
    processed_names = set()
    
    # 按优先级处理：body -> header -> footer
    for location in ["body", "header", "footer"]:
        for item in placeholders.get(location, []):
            name = item["name"]
            section_no = item.get("section_no", 1)
            if name not in processed_names:
                mapping = smart_infer_field_info(name, section_no)
                mapping["location"] = location
                mappings.append(mapping)
                processed_names.add(name)
            else:
                # 已存在，记录多位置信息
                for mapping in mappings:
                    if mapping["template_field"] == name:
                        if "also_in" not in mapping:
                            mapping["also_in"] = []
                        mapping["also_in"].append(location)
                        break
    
    # 处理checkbox（作为特殊类型）
    for cb in checkboxes:
        name = cb["name"]
        section_no = cb.get("section_no", 1)
        if name not in processed_names:
            mapping = {
                "template_field": name,
                "source_field": f"metadata.{name}",
                "type": "checkbox",
                "inferred_source": "metadata",
                "section_no": section_no
            }
            mappings.append(mapping)
            processed_names.add(name)
    
    # 按section_no排序（更直观）
    mappings.sort(key=lambda x: x.get("section_no", 1))
    
    return mappings


def main():
    parser = argparse.ArgumentParser(
        description="从Word模板中提取placeholder和checkbox",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
    python tools/extract_template_elements.py report_templates/template.docx
    python tools/extract_template_elements.py template.docx --output config/elements.json
    python tools/extract_template_elements.py template.docx --verbose
    
支持的placeholder格式:
    - {{name}}      (默认)
    - {{ name }}    (带空格)
    - 自定义: 使用 --pattern 参数，如: --pattern '\\{\\{(\\w+)\\}\\}'
        """
    )
    parser.add_argument("template", help="Word模板文件路径")
    parser.add_argument("--output", "-o", default="config/extracted_elements.json",
                       help="输出JSON文件路径 (默认: config/extracted_elements.json)")
    parser.add_argument("--generate-config", "-g", action="store_true",
                       help="同时生成report_config.json框架")
    parser.add_argument("--verbose", "-v", action="store_true",
                       help="显示详细的header/footer内容，用于调试")
    parser.add_argument("--pattern", "-p", default=None,
                       help="自定义placeholder匹配正则表达式 (默认: '{{\\s*(\\w+)\\s*}}')")
    
    args = parser.parse_args()
    
    # 读取模板
    print(f"正在读取模板: {args.template}")
    doc = Document(args.template)
    
    # 提取placeholder和checkbox
    print("正在提取placeholder...")
    placeholders = extract_placeholders(doc, pattern_str=args.pattern)
    
    print("正在提取checkbox...")
    checkboxes = extract_checkboxes(doc)
    
    # 统计信息
    total_placeholders = sum(len(v) for v in placeholders.values())
    print(f"\n提取结果:")
    print(f"  - Body中的placeholder: {len(placeholders['body'])}")
    if placeholders['body']:
        for ph in placeholders['body'][:5]:
            print(f"      * {ph['name']}")
        if len(placeholders['body']) > 5:
            print(f"      ... 还有 {len(placeholders['body'])-5} 个")
    
    print(f"  - Header中的placeholder: {len(placeholders['header'])}")
    if placeholders['header']:
        for ph in placeholders['header']:
            print(f"      * {ph['name']} (section {ph.get('section_no', 1)})")
    
    print(f"  - Footer中的placeholder: {len(placeholders['footer'])}")
    if placeholders['footer']:
        for ph in placeholders['footer']:
            print(f"      * {ph['name']} (section {ph.get('section_no', 1)})")
    
    print(f"  - Checkbox数量: {len(checkboxes)}")
    print(f"  - 总计: {total_placeholders + len(checkboxes)}")
    
    # 详细的header/footer诊断信息
    if args.verbose:
        print(f"\n文档结构诊断:")
        for i, section in enumerate(doc.sections):
            print(f"  Section {i+1}:")
            headers = [
                ('header', section.header),
                ('first_page_header', section.first_page_header),
                ('even_page_header', section.even_page_header)
            ]
            for name, header in headers:
                if header:
                    para_count = len(header.paragraphs)
                    table_count = len(header.tables)
                    has_content = any(p.text.strip() for p in header.paragraphs)
                    status = "[有内容]" if has_content else "[空]"
                    print(f"    - {name}: {status} ({para_count}段落, {table_count}表格)")
                    
                    # Verbose模式：打印实际文本内容
                    if has_content:
                        for pi, para in enumerate(header.paragraphs):
                            if para.text.strip():
                                print(f"        段落{pi}: {para.text[:80]}")
                        for ti, table in enumerate(header.tables):
                            print(f"        表格{ti}:")
                            for ri, row in enumerate(table.rows):
                                row_text = " | ".join([cell.text[:20] for cell in row.cells if cell.text.strip()])
                                if row_text:
                                    print(f"          行{ri}: {row_text[:60]}")
            
            footers = [
                ('footer', section.footer),
                ('first_page_footer', section.first_page_footer),
                ('even_page_footer', section.even_page_footer)
            ]
            for name, footer in footers:
                if footer:
                    para_count = len(footer.paragraphs)
                    table_count = len(footer.tables)
                    has_content = any(p.text.strip() for p in footer.paragraphs)
                    status = "[有内容]" if has_content else "[空]"
                    print(f"    - {name}: {status} ({para_count}段落, {table_count}表格)")
                    
                    # Verbose模式：打印实际文本内容
                    if has_content:
                        for pi, para in enumerate(footer.paragraphs):
                            if para.text.strip():
                                print(f"        段落{pi}: {para.text[:80]}")
                        for ti, table in enumerate(footer.tables):
                            print(f"        表格{ti}:")
                            for ri, row in enumerate(table.rows):
                                row_text = " | ".join([cell.text[:20] for cell in row.cells if cell.text.strip()])
                                if row_text:
                                    print(f"          行{ri}: {row_text[:60]}")
    
    # 生成输出
    output = {
        "template_path": args.template,
        "placeholders": placeholders,
        "checkboxes": checkboxes,
        "total_elements": total_placeholders + len(checkboxes)
    }
    
    # 保存提取结果
    with open(args.output, 'w', encoding='utf-8') as f:
        json.dump(output, f, ensure_ascii=False, indent=2)
    print(f"\n提取结果已保存: {args.output}")
    
    # 生成report_config框架
    if args.generate_config:
        config_output = args.output.replace('.json', '_config.json')
        field_mappings = generate_field_mappings(placeholders, checkboxes)
        
        config = {
            "template_path": args.template,
            "output_dir": "output/",
            "field_mappings": field_mappings,
            "merge_strategy": "prefer_first"
        }
        
        with open(config_output, 'w', encoding='utf-8') as f:
            json.dump(config, f, ensure_ascii=False, indent=2)
        print(f"配置框架已生成: {config_output}")
        print(f"\n提示: 请检查生成的配置，特别是:")
        print(f"  1. source_field是否正确（metadata/extracted_data/calculated_data）")
        print(f"  2. type是否正确（text/table/image/checkbox）")
        print(f"  3. 计算字段的function和args是否正确")
        print(f"  4. table类型的table_template_path是否正确")
    
    # 列出需要手动确认/修改的字段
    print("\n需要特别关注的字段:")
    for mapping in generate_field_mappings(placeholders, checkboxes):
        if mapping.get("suggested_function"):
            print(f"  [计算字段] {mapping['template_field']} -> {mapping.get('suggested_function', '需要手动指定')}")
        if mapping.get("also_in"):
            print(f"  [多位置] {mapping['template_field']} -> 主要位置: {mapping['location']}, 同时存在于: {', '.join(mapping['also_in'])}")


if __name__ == "__main__":
    main()
