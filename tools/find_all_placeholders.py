#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
查找文档中所有可能的placeholder格式
"""

import re
from docx import Document

def find_all_placeholders(doc_path: str):
    print(f"分析文档: {doc_path}")
    print("="*60)
    
    doc = Document(doc_path)
    
    # 更宽松的正则，捕获{{}}内的所有内容
    loose_pattern = re.compile(r'\{\{(.*?)\}\}')
    
    all_matches = []
    
    # 检查body段落
    for pi, para in enumerate(doc.paragraphs):
        text = para.text
        matches = loose_pattern.findall(text)
        for m in matches:
            all_matches.append(('paragraph', pi, m, text.strip()[:80]))
    
    # 检查body表格
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                text = cell.text
                matches = loose_pattern.findall(text)
                for m in matches:
                    all_matches.append(('body_table', f"{ti},{ri},{ci}", m, text.strip()[:80]))
    
    # 检查header/footer
    for section_idx, section in enumerate(doc.sections):
        # Headers
        for header_name, header in [
            ('header', section.header),
            ('first_page_header', section.first_page_header),
            ('even_page_header', section.even_page_header)
        ]:
            if header:
                for para in header.paragraphs:
                    text = para.text
                    matches = loose_pattern.findall(text)
                    for m in matches:
                        all_matches.append((header_name, section_idx, m, text.strip()[:80]))
                for ti, table in enumerate(header.tables):
                    for ri, row in enumerate(table.rows):
                        for ci, cell in enumerate(row.cells):
                            text = cell.text
                            matches = loose_pattern.findall(text)
                            for m in matches:
                                all_matches.append((f"{header_name}_table", f"sec{section_idx},t{ti},{ri},{ci}", m, text.strip()[:80]))
        
        # Footers
        for footer_name, footer in [
            ('footer', section.footer),
            ('first_page_footer', section.first_page_footer),
            ('even_page_footer', section.even_page_footer)
        ]:
            if footer:
                for para in footer.paragraphs:
                    text = para.text
                    matches = loose_pattern.findall(text)
                    for m in matches:
                        all_matches.append((footer_name, section_idx, m, text.strip()[:80]))
    
    # 去重并分类显示
    unique_names = sorted(set(m[2] for m in all_matches))
    
    print(f"\n找到 {len(all_matches)} 个占位符实例，{len(unique_names)} 个唯一名称:\n")
    
    # 按格式分类
    standard_names = [n for n in unique_names if re.match(r'^[\w.]+$', n)]
    special_names = [n for n in unique_names if not re.match(r'^[\w.]+$', n)]
    
    if standard_names:
        print("标准格式 (字母/数字/下划线/点号):")
        for name in standard_names:
            count = sum(1 for m in all_matches if m[2] == name)
            locations = set(m[0] for m in all_matches if m[2] == name)
            print(f"  - {name} (出现{count}次, 位置: {', '.join(locations)})")
    
    if special_names:
        print("\n特殊格式 (包含空格或其他字符):")
        for name in special_names:
            count = sum(1 for m in all_matches if m[2] == name)
            locations = set(m[0] for m in all_matches if m[2] == name)
            print(f"  - {repr(name)} (出现{count}次, 位置: {', '.join(locations)})")
    
    # 显示包含点号的
    dot_names = [n for n in unique_names if '.' in n]
    if dot_names:
        print(f"\n包含点号的占位符 ({len(dot_names)}个):")
        for name in dot_names:
            print(f"  - {name}")
    
    return all_matches

if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("用法: python tools/find_all_placeholders.py template.docx")
        sys.exit(1)
    
    find_all_placeholders(sys.argv[1])
