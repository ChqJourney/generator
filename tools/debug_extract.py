#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
专门调试extract_template_elements的header提取
"""

import re
from docx import Document

def debug_extraction(doc_path: str):
    print(f"分析文档: {doc_path}")
    print("="*60)
    
    doc = Document(doc_path)
    # 支持字母、数字、下划线、点号（如 {{v.15.1}}）
    pattern = re.compile(r'\{\{\s*([\w.]+)\s*\}\}')
    
    # 只检查First Page Header
    for section_idx, section in enumerate(doc.sections, 1):
        fph = section.first_page_header
        if not fph:
            print(f"Section {section_idx}: 没有First Page Header")
            continue
        
        print(f"\nSection {section_idx}: First Page Header")
        print("-" * 40)
        
        # 检查表格
        if fph.tables:
            print(f"表格数量: {len(fph.tables)}")
            for ti, table in enumerate(fph.tables, 1):
                print(f"\n  表格[{ti}]:")
                for ri, row in enumerate(table.rows, 1):
                    for ci, cell in enumerate(row.cells, 1):
                        # 方法1: cell.text
                        cell_text = cell.text or ""
                        matches1 = pattern.findall(cell_text)
                        if matches1:
                            print(f"    cell.text找到: {matches1}")
                        
                        # 方法2: 遍历paragraphs
                        para_texts = []
                        for para in cell.paragraphs:
                            pt = para.text or ""
                            if pt:
                                para_texts.append(pt)
                            matches2 = pattern.findall(pt)
                            if matches2:
                                print(f"    para.text找到: {matches2}")
                        
                        # 检查是否有差异
                        if matches1 and not any(pattern.findall(pt) for pt in para_texts):
                            print(f"    ⚠️ 只在cell.text中找到，paragraphs中没有!")
                            print(f"       cell.text = {repr(cell_text[:100])}")
                            print(f"       paragraphs = {para_texts}")
        else:
            print("没有表格")

if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("用法: python tools/debug_extract.py template.docx")
        sys.exit(1)
    
    debug_extraction(sys.argv[1])
