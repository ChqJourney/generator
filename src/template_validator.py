#!/usr/bin/env python3
"""
Template Validator Tool

检查 Word 模板中的占位符：
1. 验证占位符是否在 config 中有对应的配置
2. 检查占位符是否被分割到多个 runs 中

使用方法:
    python template_validator.py --template report_templates/production_template.docx --config config/report_config.json
"""

import argparse
import json
import re
from docx import Document
from typing import List, Dict, Set, Tuple

from utils.logging_config import get_logger

logger = get_logger(__name__)


class TemplateValidator:
    def __init__(self, template_path: str, config_path: str = None):
        self.template_path = template_path
        self.config_path = config_path
        self.doc = Document(template_path)
        self.config = self._load_config() if config_path else None
        
        # 占位符统计
        self.all_placeholders = set()
        self.split_placeholders = []  # 被分割的占位符
        self.single_run_placeholders = []  # 单个 run 的占位符
        self.undefined_placeholders = []  # 在 config 中未定义的占位符
        
    def _load_config(self) -> Dict:
        """加载配置文件（支持 JSONC 格式）"""
        try:
            from utils.jsonc_utils import load_json
            return load_json(self.config_path)
        except Exception as e:
            logger.warning(f"无法加载配置文件: {e}")
            return None
    
    def _get_configured_placeholders(self) -> Set[str]:
        """从配置中获取所有已定义的占位符"""
        if not self.config:
            return set()
        
        configured = set()
        field_mappings = self.config.get('field_mappings', [])
        
        for mapping in field_mappings:
            template_field = mapping.get('template_field')
            if template_field:
                configured.add(template_field)
        
        return configured
    
    def _extract_placeholder_from_text(self, text: str) -> List[str]:
        """从文本中提取占位符名称（去掉 {{ 和 }}）"""
        placeholders = []
        # 匹配 {{placeholder}} 格式
        pattern = r'\{\{([^}]+)\}\}'
        matches = re.findall(pattern, text)
        for match in matches:
            placeholder = match.strip()
            if placeholder:
                placeholders.append(placeholder)
        return placeholders
    
    def _check_runs_in_paragraph(self, paragraph, location: str = "paragraph") -> None:
        """检查段落中的 runs，识别被分割的占位符"""
        runs = paragraph.runs
        if not runs:
            return
        
        # 收集所有 runs 的文本
        run_texts = [run.text for run in runs]
        full_text = ''.join(run_texts)
        
        # 提取所有占位符
        placeholders_in_text = self._extract_placeholder_from_text(full_text)
        
        for placeholder in placeholders_in_text:
            self.all_placeholders.add(placeholder)
            
            # 检查这个占位符是否被分割到多个 runs
            wrapped = f"{{{{{placeholder}}}}}"
            
            # 简单情况：在一个 run 中
            single_run_count = sum(1 for rt in run_texts if wrapped in rt)
            
            if single_run_count > 0:
                self.single_run_placeholders.append({
                    'placeholder': placeholder,
                    'location': location,
                    'text': full_text[:100]
                })
            else:
                # 被分割到多个 runs
                self.split_placeholders.append({
                    'placeholder': placeholder,
                    'location': location,
                    'runs': run_texts,
                    'text': full_text[:100]
                })
    
    def validate(self) -> Dict:
        """执行完整的验证"""
        logger.info(f"开始验证模板: {self.template_path}")
        if self.config:
            logger.info(f"配置文件: {self.config_path}")
        else:
            logger.info("未提供配置文件，将只检查占位符是否被分割")
        
        # 1. 检查正文段落
        for i, para in enumerate(self.doc.paragraphs):
            if para.text.strip():
                self._check_runs_in_paragraph(para, f"paragraph {i}")
        
        # 2. 检查表格中的单元格
        for table_idx, table in enumerate(self.doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para in cell.paragraphs:
                        if para.text.strip():
                            location = f"table {table_idx}, row {row_idx}, cell {cell_idx}"
                            self._check_runs_in_paragraph(para, location)
        
        # 3. 检查是否在 config 中有定义
        if self.config:
            configured_placeholders = self._get_configured_placeholders()
            for placeholder in self.all_placeholders:
                if placeholder not in configured_placeholders:
                    self.undefined_placeholders.append(placeholder)
        
        return self._generate_report()
    
    def _generate_report(self) -> Dict:
        """生成验证报告"""
        report = {
            'total_placeholders': len(self.all_placeholders),
            'split_count': len(self.split_placeholders),
            'single_run_count': len(self.single_run_placeholders),
            'undefined_count': len(self.undefined_placeholders) if self.config else None,
        }
        
        logger.info("=" * 80)
        logger.info("验证报告")
        logger.info("=" * 80)
        logger.info(f"发现的占位符总数: {report['total_placeholders']}")
        logger.info(f"  - 单个 run 中的占位符: {report['single_run_count']}")
        logger.info(f"  - 被分割到多个 runs 的占位符: {report['split_count']}")

        if self.config:
            logger.info(f"  - 在配置中未定义的占位符: {report['undefined_count']}")

        # 显示被分割的占位符详情
        if self.split_placeholders:
            logger.info("-" * 80)
            logger.info("被分割到多个 runs 的占位符（需要修复）：")
            logger.info("-" * 80)
            for item in self.split_placeholders:
                logger.info(f"  占位符: {{{{{item['placeholder']}}}}}")
                logger.info(f"  位置: {item['location']}")
                logger.info(f"  文本: {item['text'][:80]}...")
                logger.info(f"  Runs 详情:")
                for i, run_text in enumerate(item['runs']):
                    logger.info(f"    Run {i}: '{run_text}'")

        # 显示未定义的占位符
        if self.undefined_placeholders:
            logger.info("-" * 80)
            logger.info("在配置中未定义的占位符：")
            logger.info("-" * 80)
            for placeholder in self.undefined_placeholders:
                logger.info(f"  - {{{{{placeholder}}}}}")

        # 显示建议
        logger.info("=" * 80)
        logger.info("建议")
        logger.info("=" * 80)

        if self.split_placeholders:
            logger.info("1. 修复被分割的占位符：")
            logger.info("   这些占位符在 Word 中被分割成了多个 runs，导致 {{}} 无法被正确清除。")
            logger.info("   建议：在 Word 中重新输入这些占位符，确保它们在一个连续的文本块中。")

        if self.undefined_placeholders:
            logger.info("2. 添加缺失的占位符配置：")
            logger.info(f"   在配置文件 '{self.config_path}' 中添加这些占位符的 field_mappings。")

        if not self.split_placeholders and not self.undefined_placeholders:
            logger.info("模板验证通过！所有占位符都在单个 run 中，并且都已在配置中定义。")
        
        return report


def main():
    parser = argparse.ArgumentParser(
        description='验证 Word 模板中的占位符',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
    # 只检查占位符是否被分割
    python template_validator.py --template report_templates/production_template.docx
    
    # 同时检查占位符是否在 config 中有定义
    python template_validator.py --template report_templates/production_template.docx --config config/report_config.json
        """
    )
    
    parser.add_argument(
        '--template', '-t',
        required=True,
        help='Word 模板文件路径 (.docx)'
    )
    
    parser.add_argument(
        '--config', '-c',
        help='配置文件路径 (.json)，可选'
    )
    
    args = parser.parse_args()
    
    # 执行验证
    validator = TemplateValidator(args.template, args.config)
    report = validator.validate()
    
    # 返回退出码
    exit_code = 0
    if report['split_count'] > 0:
        exit_code = 1
    if args.config and report['undefined_count'] > 0:
        exit_code = 1
    
    return exit_code


if __name__ == '__main__':
    exit(main())
