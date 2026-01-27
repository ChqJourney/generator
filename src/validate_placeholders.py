"""
Validate placeholder in Word documents
"""
import sys
import json
import re
from pathlib import Path
from docx import Document
from typing import Set, Dict, List, Tuple

def extract_placeholders_from_paragraphs(paragraphs: List) -> Set[Tuple[str, str]]:
    placeholders = set()
    placeholder_pattern = re.compile(r'\{\{(\w+)\}\}')
    
    for para in paragraphs:
        full_text = para.text
        matches = placeholder_pattern.findall(full_text)
        for match in matches:
            placeholders.add((match, 'paragraph'))
    
    return placeholders

def extract_placeholders_from_table(table) -> Set[Tuple[str, str]]:
    placeholders = set()
    placeholder_pattern = re.compile(r'\{\{(\w+)\}\}')
    
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                full_text = para.text
                matches = placeholder_pattern.findall(full_text)
                for match in matches:
                    placeholders.add((match, 'table'))
    
    return placeholders

def extract_all_placeholders(doc: Document) -> Dict[str, any]:
    result = {
        'body': set(),
        'headers': {},
        'footers': {}
    }
    
    result['body'].update(extract_placeholders_from_paragraphs(doc.paragraphs))
    
    for table in doc.tables:
        result['body'].update(extract_placeholders_from_table(table))
    
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header:
                header_placeholders = extract_placeholders_from_paragraphs(header.paragraphs)
                for table in header.tables:
                    header_placeholders.update(extract_placeholders_from_table(table))
                if header_placeholders:
                    result['headers'][str(header)] = header_placeholders
        
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer:
                footer_placeholders = extract_placeholders_from_paragraphs(footer.paragraphs)
                for table in footer.tables:
                    footer_placeholders.update(extract_placeholders_from_table(table))
                if footer_placeholders:
                    result['footers'][str(footer)] = footer_placeholders
    
    return result

def load_field_mappings(config_path: Path) -> List[Dict]:
    with open(config_path, 'r', encoding='utf-8') as f:
        config = json.load(f)
    return config.get('field_mappings', [])

def validate_placeholders_in_paragraphs(all_placeholders: Dict[str, any]) -> List[str]:
    errors = []
    
    total_placeholders = 0
    paragraph_placeholders = 0
    
    for location, placeholders in all_placeholders.items():
        if location == 'headers':
            for header_key, header_placeholders in placeholders.items():
                for name, loc in header_placeholders:
                    total_placeholders += 1
                    if loc == 'paragraph':
                        paragraph_placeholders += 1
                    else:
                        errors.append(f"Placeholder '{name}' in header is in a table, not a paragraph")
        elif location == 'footers':
            for footer_key, footer_placeholders in placeholders.items():
                for name, loc in footer_placeholders:
                    total_placeholders += 1
                    if loc == 'paragraph':
                        paragraph_placeholders += 1
                    else:
                        errors.append(f"Placeholder '{name}' in footer is in a table, not a paragraph")
        else:
            for name, loc in placeholders:
                total_placeholders += 1
                if loc == 'paragraph':
                    paragraph_placeholders += 1
                else:
                    errors.append(f"Placeholder '{name}' in body is a table placeholder")
    
    if total_placeholders > 0 and paragraph_placeholders == total_placeholders:
        errors.append(f"Validation failed: All {total_placeholders} placeholders are in paragraphs (expected some in tables for table placeholders)")
    
    return errors

def validate_field_mappings_exist(field_mappings: List[Dict], template_path: Path) -> List[str]:
    doc = Document(template_path)
    all_placeholders = extract_all_placeholders(doc)
    
    template_placeholders = set()
    for location, placeholders in all_placeholders.items():
        if location == 'headers':
            for header_placeholders in placeholders.values():
                template_placeholders.update({name for name, _ in header_placeholders})
        elif location == 'footers':
            for footer_placeholders in placeholders.values():
                template_placeholders.update({name for name, _ in footer_placeholders})
        else:
            template_placeholders.update({name for name, _ in placeholders})
    
    errors = []
    for mapping in field_mappings:
        template_field = mapping.get('template_field')
        if template_field not in template_placeholders:
            errors.append(f"Field mapping 'template_field' '{template_field}' not found in template: {template_path}")
    
    return errors

def validate_extra_placeholders(field_mappings: List[Dict], template_path: Path) -> List[str]:
    doc = Document(template_path)
    all_placeholders = extract_all_placeholders(doc)
    
    template_placeholders = set()
    for location, placeholders in all_placeholders.items():
        if location == 'headers':
            for header_placeholders in placeholders.values():
                template_placeholders.update({name for name, _ in header_placeholders})
        elif location == 'footers':
            for footer_placeholders in placeholders.values():
                template_placeholders.update({name for name, _ in footer_placeholders})
        else:
            template_placeholders.update({name for name, _ in placeholders})
    
    config_placeholders = {mapping.get('template_field') for mapping in field_mappings}
    
    errors = []
    extra_placeholders = template_placeholders - config_placeholders
    for placeholder in extra_placeholders:
        errors.append(f"Placeholder '{placeholder}' found in template but not in field_mappings")
    
    return errors

def main():
    import argparse
    parser = argparse.ArgumentParser(description='Validate placeholder in Word documents')
    parser.add_argument('--config', required=True, help='Path to config file (e.g., config/report_config.json)')
    parser.add_argument('--template', help='Path to template file (overrides config)')
    args = parser.parse_args()
    
    try:
        config_path = Path(args.config)
        
        if not config_path.exists():
            print(f"Error: Config file not found: {config_path}", file=sys.stderr)
            return 1
        
        field_mappings = load_field_mappings(config_path)
        
        with open(config_path, 'r', encoding='utf-8') as f:
            config = json.load(f)
        
        template_path = Path(args.template) if args.template else Path('report_templates/report_template1.docx')
        if not template_path.is_absolute():
            template_path = config_path.parent.parent / template_path
        
        if not template_path.exists():
            print(f"Error: Template file not found: {template_path}", file=sys.stderr)
            return 1
        
        print(f"Validating template: {template_path}")
        print(f"Config file: {config_path}")
        print()
        
        doc = Document(str(template_path))
        all_placeholders = extract_all_placeholders(doc)
        
        print("=== Placeholder Locations ===")
        for location, placeholders in all_placeholders.items():
            if location == 'headers':
                for i, (header_key, header_placeholders) in enumerate(placeholders.items(), 1):
                    print(f"\nHeader {i}:")
                    for name, loc in sorted(header_placeholders):
                        print(f"  - {{{{{name}}}}} in {loc}")
            elif location == 'footers':
                for i, (footer_key, footer_placeholders) in enumerate(placeholders.items(), 1):
                    print(f"\nFooter {i}:")
                    for name, loc in sorted(footer_placeholders):
                        print(f"  - {{{{{name}}}}} in {loc}")
            else:
                print(f"\n{location.capitalize()}:")
                for name, loc in sorted(placeholders):
                    print(f"  - {{{{{name}}}}} in {loc}")
        
        print("\n=== Validation Results ===")
        
        has_errors = False
        
        errors1 = validate_placeholders_in_paragraphs(all_placeholders)
        if errors1:
            has_errors = True
            print("\n[X] Validation 1: Check if placeholders are in paragraphs")
            for error in errors1:
                print(f"  - {error}")
        else:
            print("\n[OK] Validation 1: All placeholders are correctly located")
        
        errors2 = validate_field_mappings_exist(field_mappings, template_path)
        if errors2:
            has_errors = True
            print("\n[X] Validation 2: Check if field_mappings exist in template")
            for error in errors2:
                print(f"  - {error}")
        else:
            print("\n[OK] Validation 2: All field_mappings exist in template")
        
        errors3 = validate_extra_placeholders(field_mappings, template_path)
        if errors3:
            has_errors = True
            print("\n[X] Validation 3: Check for extra placeholders not in field_mappings")
            for error in errors3:
                print(f"  - {error}")
        else:
            print("\n[OK] Validation 3: No extra placeholders found")
        
        print(f"\n{'='*50}")
        if has_errors:
            print("VALIDATION FAILED")
            return 1
        else:
            print("VALIDATION PASSED")
            return 0
    
    except FileNotFoundError as e:
        print(f"Error: File not found - {e}", file=sys.stderr)
        return 1
    except json.JSONDecodeError as e:
        print(f"Error: Invalid JSON in config file - {e}", file=sys.stderr)
        return 1
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        import traceback
        traceback.print_exc()
        return 1

if __name__ == '__main__':
    sys.exit(main())