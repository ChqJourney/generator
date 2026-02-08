from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import shutil
from typing import List, Dict, Optional, Tuple, Any
from abc import ABC, abstractmethod
from copy import deepcopy
from table_processor import TableDataTransformer
from utils.logging_config import get_logger

logger = get_logger(__name__)

class DocxTemplateError(Exception):
    pass

class PlaceholderNotFoundError(DocxTemplateError):
    def __init__(self, placeholder: str, location: str = 'body'):
        self.placeholder = placeholder
        self.location = location
        super().__init__(f"Placeholder '{placeholder}' not found in {location}")

class InvalidLocationError(DocxTemplateError):
    def __init__(self, location: str, content_type: str):
        super().__init__(f"Invalid location '{location}' for {content_type}. Valid locations: body, header, footer")

class PlaceholderFinder:
    @staticmethod
    def _search_paragraphs_in_container(container, placeholder):
        # 同时搜索裸占位符和带 {{}} 的占位符
        search_terms = [placeholder, f'{{{{{placeholder}}}}}']
        for i, paragraph in enumerate(container.paragraphs):
            if any(term in paragraph.text for term in search_terms):
                yield i, paragraph
        for table_idx, table in enumerate(container.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for p_idx, paragraph in enumerate(cell.paragraphs):
                        if any(term in paragraph.text for term in search_terms):
                            yield (table_idx, row_idx, cell_idx, p_idx), paragraph

    @staticmethod
    def find_all_placeholders_in_location(doc: Document, placeholder: str, location: str = 'body') -> List[Tuple[Any, Any]]:
        results = []
        if location == 'body':
            results.extend(list(PlaceholderFinder._search_paragraphs_in_container(doc, placeholder)))
        elif location == 'header':
            for section in doc.sections:
                for header in [section.header, section.first_page_header, section.even_page_header]:
                    if header:
                        results.extend(list(PlaceholderFinder._search_paragraphs_in_container(header, placeholder)))
        elif location == 'footer':
            for section in doc.sections:
                for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                    if footer:
                        results.extend(list(PlaceholderFinder._search_paragraphs_in_container(footer, placeholder)))
        return results

    @staticmethod
    def find_paragraph_with_placeholder(doc: Document, placeholder: str, location: str = 'body') -> Tuple[Any, Any]:
        results = PlaceholderFinder.find_all_placeholders_in_location(doc, placeholder, location)
        return results[0] if results else (None, None)

    @staticmethod
    def replace_paragraph_with_element(paragraph, element, location: str = 'body'):
        try:
            p_element = paragraph._element
            p_parent = p_element.getparent()
            
            if p_parent is None:
                raise DocxTemplateError("Cannot find parent element of paragraph")
            
            # 获取父容器中的索引位置
            try:
                index = list(p_parent).index(p_element)
            except ValueError:
                # 如果在父容器中找不到，可能是因为paragraph在特殊容器中
                raise DocxTemplateError(f"Paragraph not found in parent container")
            
            # 移除原段落
            p_parent.remove(p_element)
            
            # 插入新元素
            # 注意：对于表格元素，需要使用深拷贝以避免重复引用
            new_element = deepcopy(element)
            p_parent.insert(index, new_element)
            
        except (AttributeError, ValueError) as e:
            raise DocxTemplateError(f"Failed to replace paragraph with element: {str(e)}")

    @staticmethod
    def _replace_in_paragraph(paragraph, placeholder, value):
        # 只搜索带 {{}} 的完整占位符格式，避免部分匹配问题（如 Pon 匹配到 Ponmax）
        wrapped_placeholder = f'{{{{{placeholder}}}}}'
        
        # 检查段落中是否包含完整格式的占位符
        if wrapped_placeholder not in paragraph.text:
            return False
        
        runs = paragraph.runs
        if not runs:
            return False
        
        # 首先尝试在单个 run 中匹配完整格式 {{placeholder}}
        for run in runs:
            if wrapped_placeholder in run.text:
                run.text = run.text.replace(wrapped_placeholder, value)
                return True
        
        # 处理占位符被分割到多个 runs 的情况
        # 例如: Run0='{{', Run1='efficacy', Run2='}}'
        full_text = ''.join(run.text for run in runs)
        if wrapped_placeholder in full_text:
            new_text = full_text.replace(wrapped_placeholder, value)
            # 将所有文本放入第一个 run，清空其他 runs
            runs[0].text = new_text
            for run in runs[1:]:
                run.text = ''
            return True
        
        # 处理只有占位符名称被分割，但 {{ 和 }} 分别在前后 run 的情况
        # 例如: Run0='{{', Run1='efficacy', Run2='}}'
        # 或者值已经被替换的情况: Run0='{{', Run1='100.00', Run2='}}'
        for i, run in enumerate(runs):
            # 检查是否包含占位符名称或者已经被替换的值
            has_placeholder = placeholder in run.text
            has_value = value in run.text and value != placeholder
            
            if not has_placeholder and not has_value:
                continue
            
            # 检查前一个 run 是否包含 '{{'
            has_open_braces = False
            if i > 0:
                prev_text = runs[i-1].text.strip()
                # 支持 '{{' 或末尾有 '{{'
                if prev_text == '{{' or prev_text.endswith('{{'):
                    has_open_braces = True
            
            # 检查后一个 run 是否包含 '}}'
            has_close_braces = False
            if i < len(runs) - 1:
                next_text = runs[i+1].text.strip()
                # 支持 '}}' 或开头有 '}}'
                if next_text == '}}' or next_text.startswith('}}'):
                    has_close_braces = True
            
            # 如果包含占位符，进行替换
            if has_placeholder:
                run.text = run.text.replace(placeholder, value)
            
            # 清除前一个 run 的 '{{'
            if has_open_braces:
                prev_text = runs[i-1].text
                if prev_text.strip() == '{{':
                    runs[i-1].text = ''
                else:
                    # 只移除末尾的 '{{'
                    idx = prev_text.rfind('{{')
                    if idx >= 0:
                        runs[i-1].text = prev_text[:idx]
            
            # 清除后一个 run 的 '}}'
            if has_close_braces:
                next_text = runs[i+1].text
                if next_text.strip() == '}}':
                    runs[i+1].text = ''
                else:
                    # 只移除开头的 '}}'
                    idx = next_text.find('}}')
                    if idx >= 0:
                        runs[i+1].text = next_text[idx+2:]
            
            return True
        
        return False

class ContentInserter(ABC):
    def __init__(self, doc: Document):
        self.doc = doc
    
    @abstractmethod
    def insert(self, *args, **kwargs):
        pass
    
    def validate_location(self, location: str, valid_locations: List[str]):
        if location not in valid_locations:
            raise InvalidLocationError(location, self.__class__.__name__)

class TextInserter(ContentInserter):
    def insert(self, placeholder: str, value: str, location: str = 'body'):
        self.validate_location(location, ['body', 'header', 'footer'])
        
        # 如果没有指定location或者在指定location找不到，则在所有位置查找
        results = PlaceholderFinder.find_all_placeholders_in_location(self.doc, placeholder, location)
        
        # 如果在指定位置找不到，尝试在所有位置查找
        if not results:
            logger.warning(f"在 {location} 中未找到占位符 '{placeholder}'，尝试在所有位置查找...")
            for loc in ['header', 'body', 'footer']:
                if loc != location:  # 跳过已经搜索过的位置
                    results = PlaceholderFinder.find_all_placeholders_in_location(self.doc, placeholder, loc)
                    if results:
                        logger.info(f"在 {loc} 中找到占位符 '{placeholder}'")
                        location = loc  # 更新location
                        break
        
        if not results:
            logger.warning(f"占位符 '{placeholder}' 在所有位置都未找到，跳过此操作")
            return
        
        replaced = False
        for idx, paragraph in results:
            if PlaceholderFinder._replace_in_paragraph(paragraph, placeholder, value):
                replaced = True
        
        if not replaced:
            print(f"警告: 占位符 '{placeholder}' 未能成功替换，跳过此操作")

class TableInserter(ContentInserter):
    def insert(self, placeholder: str, table_template_path: str, 
               raw_data: Optional[List[List[str]]] = None,
               transformations: Optional[List[Dict]] = None,
               calculated_report: Optional[Dict] = None,
               row_strategy: str = 'fixed_rows',
               skip_columns: Optional[List[int]] = None,
               header_rows: int = 1,
               text_insert: Optional[List[Dict]] = None,
               location: str = 'body'):
        self.validate_location(location, ['body'])
        
        if not os.path.exists(table_template_path):
            raise DocxTemplateError(f"Table template file not found: {table_template_path}")
        
        transformer = TableDataTransformer()
        
        processed_data = raw_data
        if raw_data and transformations:
            processed_data = transformer.transform(raw_data, transformations, calculated_report)
        
        table_template = Document(table_template_path)
        if not table_template.tables:
            raise DocxTemplateError(f"No tables found in template file: {table_template_path}")
        
        template_table = table_template.tables[0]
        
        if row_strategy == 'fixed_rows':
            print(f"processed data: {processed_data}")
            self._fill_fixed_rows(template_table, processed_data, skip_columns, header_rows, text_insert, calculated_report)
        elif row_strategy == 'dynamic_rows':
            self._fill_dynamic_rows(template_table, processed_data, skip_columns, header_rows)
        
        results = PlaceholderFinder.find_all_placeholders_in_location(self.doc, placeholder, location)
        
        if not results:
            print(f"警告: 在 {location} 中未找到占位符 '{placeholder}'，尝试在 body 中查找...")
            if location != 'body':
                results = PlaceholderFinder.find_all_placeholders_in_location(self.doc, placeholder, 'body')
                if results:
                    print(f"在 body 中找到占位符 '{placeholder}'")
                    location = 'body'
        
        if not results:
            print(f"警告: 占位符 '{placeholder}' 未找到，跳过表格插入操作")
            return
        
        for idx, paragraph in results:
            try:
                PlaceholderFinder.replace_paragraph_with_element(paragraph, template_table._element)
            except (AttributeError, ValueError, TypeError, DocxTemplateError) as e:
                try:
                    print(f"尝试在单元格内插入表格 '{placeholder}'...")
                    self._insert_table_in_cell(paragraph, template_table)
                except Exception as e2:
                    print(f"警告: 无法替换占位符 '{placeholder}' 为表格: {str(e)}")
                    print(f"      尝试在单元格内插入也失败: {str(e2)}")
                    continue
    
    def _fill_fixed_rows(self, table: Any, data: List[List[Any]], skip_columns: Optional[List[int]], header_rows: int, text_insert: Optional[List[Dict]] = None, calculated_report: Optional[Dict] = None):
        """填充固定行数的表格
        
        header_rows: 模板表格中表头行数，从第 header_rows 行开始填充数据
        data: 要填充的完整数据（不包含表头）
        text_insert: 在特定行列插入文本的配置列表
        calculated_report: 用于解析 value 路径的数据源
        """
        if not data:
            return
        
        # header_rows 只控制从模板表格的哪一行开始填充，不跳过数据
        data_row_idx = 0
        for row_idx, row in enumerate(table.rows):
            # 跳过模板表格的表头行
            if row_idx < header_rows:
                continue
            
            # 使用完整的数据，不跳过任何行
            if data_row_idx >= len(data):
                break
            
            data_row = data[data_row_idx]
            data_col_idx = 0
            
            for col_idx, cell in enumerate(row.cells):
                if skip_columns and col_idx in skip_columns:
                    continue
                
                if data_col_idx < len(data_row):
                    value = data_row[data_col_idx]
                    if value is None or value == '':
                        pass
                    else:
                        self._set_cell_value(cell, str(value))
                    data_col_idx += 1
            
            data_row_idx += 1
        
        # 处理 text_insert - 在特定行列插入文本
        if text_insert and calculated_report:
            for insert_config in text_insert:
                row_idx = insert_config.get('row')
                col_idx = insert_config.get('column')
                value_path = insert_config.get('value')
                
                if row_idx is not None and col_idx is not None and value_path:
                    # 从 calculated_report 中解析值
                    value = self._get_value_from_path(calculated_report, value_path)
                    if value:
                        # 获取指定行列的 cell 并设置值
                        if row_idx < len(table.rows):
                            row = table.rows[row_idx]
                            if col_idx < len(row.cells):
                                cell = row.cells[col_idx]
                                self._set_cell_value(cell, str(value))
    
    def _get_value_from_path(self, data: Dict, path: str) -> Optional[str]:
        """从字典中根据路径获取值
        
        Args:
            data: 数据源字典
            path: 点分隔的路径，如 'extracted_data.light_source_type'
            
        Returns:
            找到的值或 None
        """
        if not path or not data:
            return None
        
        keys = path.split('.')
        current = data
        
        for key in keys:
            if isinstance(current, dict) and key in current:
                current = current[key]
            else:
                return None
        
        return str(current) if current is not None else None
    
    def _fill_dynamic_rows(self, table: Any, data: List[List[Any]], skip_columns: Optional[List[int]], header_rows: int):
        """动态填充表格行数
        
        header_rows: 模板表格中表头行数，从第 header_rows 行开始填充数据
        data: 要填充的完整数据（不包含表头）
        """
        if not data:
            return
        
        # 保留表头行，删除多余的数据行
        while len(table.rows) > header_rows:
            table._tbl.remove(table.rows[-1]._tr)
        
        num_columns = len(table.rows[0].cells) if table.rows else 0
        
        # header_rows 只控制从模板表格的哪一行开始填充，不跳过数据
        for data_row in data:
            new_row = table.add_row()
            
            if len(new_row.cells) < num_columns:
                for _ in range(num_columns - len(new_row.cells)):
                    new_row.add_cell()
            
            data_col_idx = 0
            for col_idx, cell in enumerate(new_row.cells):
                if skip_columns and col_idx in skip_columns:
                    continue
                
                if data_col_idx < len(data_row):
                    value = data_row[data_col_idx]
                    self._set_cell_value(cell, str(value) if value else '')
                    data_col_idx += 1
    
    def _set_cell_value(self, cell: Any, value: str):
        """设置单元格值"""
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.text = value
                return          # 找到run并设置后直接返回
        # 没有run则使用第一个paragraph
        if cell.paragraphs:
            cell.paragraphs[0].add_run(value)
        else:
            cell.add_paragraph(value)
    
    def _insert_table_in_cell(self, paragraph, template_table):
        cell = self._find_parent_cell(paragraph)
        if cell is None:
            raise DocxTemplateError("Cannot find parent cell for paragraph")
        
        paragraph.clear()
        cell_element = cell._element
        new_table_element = deepcopy(template_table._element)
        p_element = paragraph._element
        p_index = list(cell_element).index(p_element)
        cell_element.remove(p_element)
        cell_element.insert(p_index, new_table_element)
        print(f"成功在单元格内插入表格")
    
    def _find_parent_cell(self, paragraph):
        try:
            p_element = paragraph._element
            current = p_element.getparent()
            
            while current is not None:
                if current.tag.endswith('tc'):
                    from docx.table import _Cell
                    return _Cell(current, None)
                current = current.getparent()
            
            return None
        except Exception:
            return None

class ImageInserter(ContentInserter):
    def insert(self, placeholder: str, image_paths: List[str], width: Optional[Any] = None, 
               height: Optional[Any] = None, alignment: Optional[str] = None, location: str = 'body'):
        self.validate_location(location, ['body', 'header', 'footer'])
        
        if not image_paths:
            raise DocxTemplateError(f"No image paths provided for placeholder '{placeholder}'")
        
        # 处理图片路径，自动添加 data_files 前缀
        processed_image_paths = []
        for img_path in image_paths:
            processed_path = self._resolve_image_path(img_path)
            if not os.path.exists(processed_path):
                raise DocxTemplateError(f"Image file not found: {processed_path} (original: {img_path})")
            processed_image_paths.append(processed_path)
        
        # 使用处理后的路径
        image_paths = processed_image_paths
        
        self._validate_image_dimensions(width, height)
        
        def create_image_paragraphs(doc, parent_element=None):
            paragraphs = []
            for idx, img_path in enumerate(image_paths):
                if parent_element is None:
                    new_p = doc.add_paragraph()
                else:
                    new_p = parent_element.add_paragraph()
                
                if alignment == 'center':
                    new_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                elif alignment == 'right':
                    new_p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                elif alignment == 'left':
                    new_p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                
                run = new_p.add_run()
                try:
                    if width and height:
                        run.add_picture(img_path, width=width, height=height)
                    elif width:
                        run.add_picture(img_path, width=width)
                    elif height:
                        run.add_picture(img_path, height=height)
                    else:
                        run.add_picture(img_path, width=Inches(4.0))
                except (ValueError, TypeError, OSError) as e:
                    raise DocxTemplateError(f"Failed to insert image '{img_path}': {str(e)}")
                
                paragraphs.append(new_p)
            return paragraphs
        
        results = PlaceholderFinder.find_all_placeholders_in_location(self.doc, placeholder, location)
        
        # 如果在指定位置找不到，尝试在所有位置查找
        if not results:
            print(f"警告: 在 {location} 中未找到占位符 '{placeholder}'，尝试在所有位置查找...")
            for loc in ['header', 'body', 'footer']:
                if loc != location:
                    results = PlaceholderFinder.find_all_placeholders_in_location(self.doc, placeholder, loc)
                    if results:
                        print(f"在 {loc} 中找到占位符 '{placeholder}'")
                        location = loc
                        break
        
        if not results:
            print(f"警告: 占位符 '{placeholder}' 在所有位置都未找到，跳过图片插入操作")
            return
        
        for idx, paragraph in results:
            try:
                parent_element = self._get_parent_element(paragraph)
                self._replace_placeholder_with_images(paragraph, create_image_paragraphs, parent_element)
            except (AttributeError, ValueError, TypeError) as e:
                raise DocxTemplateError(f"Failed to replace placeholder '{placeholder}': {str(e)}")

    def _replace_placeholder_with_images(self, paragraph, create_fn, parent_element):
        p_element = paragraph._element
        p_parent = p_element.getparent()
        if p_parent is None:
            raise DocxTemplateError("Cannot find parent element of paragraph")
        
        index = list(p_parent).index(p_element)
        p_parent.remove(p_element)
        
        paragraphs = create_fn(self.doc, parent_element)
        for idx, new_p in enumerate(paragraphs):
            if idx == 0:
                p_parent.insert(index, new_p._element)
            else:
                if parent_element is None:
                    empty_p = self.doc.add_paragraph()
                else:
                    empty_p = parent_element.add_paragraph()
                p_parent.insert(index + idx * 2 - 1, empty_p._element)
                p_parent.insert(index + idx * 2, new_p._element)

    def _get_parent_element(self, paragraph):
        try:
            p_element = paragraph._element
            p_parent = p_element.getparent()
            if p_parent is None:
                return None
            for table in self.doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell._element == p_parent:
                            return cell
            return None
        except Exception:
            return None

    def _resolve_image_path(self, img_path: str) -> str:
        """解析图片路径，如果是相对路径则添加 data_files 前缀"""
        # 如果路径已经存在，直接返回
        if os.path.exists(img_path):
            return img_path
        
        # 如果是相对路径（以 ./ 或直接文件名开头），尝试添加 data_files 前缀
        if not os.path.isabs(img_path):
            # 移除开头的 ./
            clean_path = img_path.lstrip('./')
            clean_path = clean_path.lstrip('.\\')
            
            # 尝试在 data_files 目录中查找
            data_files_path = os.path.join('data_files', clean_path)
            if os.path.exists(data_files_path):
                return data_files_path
            
            # 尝试相对于当前工作目录的 data_files
            cwd_data_files_path = os.path.join(os.getcwd(), 'data_files', clean_path)
            if os.path.exists(cwd_data_files_path):
                return cwd_data_files_path
        
        # 如果都找不到，返回原路径（会在后续检查中报错）
        return img_path

    def _validate_image_dimensions(self, width, height):
        from docx.shared import Length
        valid_dimensions = []
        if width is not None:
            valid_dimensions.append(width)
        if height is not None:
            valid_dimensions.append(height)
        
        for dimension in valid_dimensions:
            if not isinstance(dimension, Length):
                raise DocxTemplateError(f"Invalid image dimension '{dimension}'. Must be a Length object (e.g., Inches, Mm, Cm, Pt)")

class CheckboxInserter(ContentInserter):
    """
    Checkbox 状态更新器
    
    根据 checkbox_mapping 更新 Word 文档中的 checkbox 控件状态
    """
    
    def insert(self, checkbox_mapping: Dict[str, bool]):
        """
        更新文档中的 checkbox 状态
        
        Args:
            checkbox_mapping: checkbox 名称到布尔值的映射，如 {"cb1": True, "cb2": False}
        """
        from docx.oxml import parse_xml
        
        if not checkbox_mapping:
            logger.warning("No checkbox mapping provided, skipping checkbox update")
            return
        
        root = self.doc.part.element
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        w_ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        
        checkboxes = root.findall('.//w:checkBox', namespaces=ns)
        updated = {}
        
        for checkbox in checkboxes:
            ffdata = checkbox.getparent()
            if ffdata is not None:
                name = ffdata.find('w:name', namespaces=ns)
                if name is not None:
                    field_name = name.get(w_ns + 'val')
                    
                    if field_name in checkbox_mapping:
                        should_check = checkbox_mapping[field_name]
                        
                        checked = checkbox.find('w:checked', namespaces=ns)
                        default = checkbox.find('w:default', namespaces=ns)
                        
                        if should_check:
                            if checked is None:
                                new_checked = parse_xml(f'<w:checked w:val="1" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
                                checkbox.append(new_checked)
                            else:
                                checked.set(w_ns + 'val', '1')
                            if default is not None:
                                default.set(w_ns + 'val', '1')
                            updated[field_name] = True
                            logger.info(f"Checkbox '{field_name}' checked")
                        else:
                            if checked is not None:
                                checkbox.remove(checked)
                            if default is not None:
                                default.set(w_ns + 'val', '0')
                            updated[field_name] = False
                            logger.info(f"Checkbox '{field_name}' unchecked")
        
        # 检查是否有未找到的 checkbox
        for field_name in checkbox_mapping:
            if field_name not in updated:
                logger.warning(f"Checkbox '{field_name}' not found in document")
        
        logger.info(f"Updated {len(updated)} checkbox(es)")

class DocxTemplateProcessor:
    def __init__(self, template_path: str, output_path: str):
        if not os.path.exists(template_path):
            raise DocxTemplateError(f"Template file not found: {template_path}")
        
        self.template_path = template_path
        self.output_path = output_path
        self.operations = []
        
        shutil.copy(template_path, output_path)
        self.doc = Document(output_path)
    
    def add_text(self, placeholder: str, value: str, location: str = 'body'):
        self.operations.append({
            'type': 'text',
            'placeholder': placeholder,
            'value': value,
            'location': location
        })
        return self
    
    def add_table(self, placeholder: str, table_template_path: str, 
                  raw_data: Optional[List[List[str]]] = None,
                  transformations: Optional[List[Dict]] = None,
                  calculated_report: Optional[Dict] = None,
                  row_strategy: str = 'fixed_rows',
                  skip_columns: Optional[List[int]] = None,
                  header_rows: int = 1,
                  text_insert: Optional[List[Dict]] = None):
        self.operations.append({
            'type': 'table',
            'placeholder': placeholder,
            'table_template_path': table_template_path,
            'raw_data': raw_data,
            'transformations': transformations,
            'calculated_report': calculated_report,
            'row_strategy': row_strategy,
            'skip_columns': skip_columns,
            'header_rows': header_rows,
            'text_insert': text_insert
        })
        return self
    
    def add_image(self, placeholder: str, image_paths: List[str], 
                  width: Optional[any] = None, height: Optional[any] = None, 
                  alignment: Optional[str] = None, location: str = 'body'):
        self.operations.append({
            'type': 'image',
            'placeholder': placeholder,
            'image_paths': image_paths,
            'width': width,
            'height': height,
            'alignment': alignment,
            'location': location
        })
        return self
    
    def add_checkboxes(self, checkbox_mapping: Dict[str, bool]):
        """
        添加 checkbox 状态更新操作
        
        Args:
            checkbox_mapping: checkbox 名称到布尔值的映射
                           如 {"cb1": True, "cb2": False}
        """
        self.operations.append({
            'type': 'checkbox',
            'checkbox_mapping': checkbox_mapping
        })
        return self
    
    def get_all_placeholders(self) -> List[str]:
        """
        获取模板中所有的占位符名称
        
        Returns:
            List[str]: 占位符名称列表
        """
        import re
        pattern = re.compile(r'\{\{([\w.]+)\}\}')
        result = []
        
        for para in self.doc.paragraphs:
            matches = pattern.findall(para.text)
            result.extend(matches)
        
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        matches = pattern.findall(para.text)
                        result.extend(matches)
        
        for section in self.doc.sections:
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header:
                    for para in header.paragraphs:
                        matches = pattern.findall(para.text)
                        result.extend(matches)
                    for table in header.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for para in cell.paragraphs:
                                    matches = pattern.findall(para.text)
                                    result.extend(matches)
            
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer:
                    for para in footer.paragraphs:
                        matches = pattern.findall(para.text)
                        result.extend(matches)
                    for table in footer.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for para in cell.paragraphs:
                                    matches = pattern.findall(para.text)
                                    result.extend(matches)
        
        return list(dict.fromkeys(result))
    
    def process(self):
        try:
            for op in self.operations:
                if op['type'] == 'text':
                    inserter = TextInserter(self.doc)
                    inserter.insert(op['placeholder'], op['value'], op['location'])
                
                elif op['type'] == 'table':
                    inserter = TableInserter(self.doc)
                    inserter.insert(
                        op['placeholder'], 
                        op['table_template_path'],
                        op.get('raw_data'),
                        op.get('transformations'),
                        op.get('calculated_report'),
                        op.get('row_strategy', 'fixed_rows'),
                        op.get('skip_columns'),
                        op.get('header_rows', 1),
                        op.get('text_insert'),
                        'body'
                    )
                
                elif op['type'] == 'image':
                    inserter = ImageInserter(self.doc)
                    inserter.insert(op['placeholder'], op['image_paths'], 
                                  op['width'], op['height'], op['alignment'], op['location'])
                
                elif op['type'] == 'checkbox':
                    inserter = CheckboxInserter(self.doc)
                    inserter.insert(op['checkbox_mapping'])
            
            self.doc.save(self.output_path)
            print(f"文档已保存至: {self.output_path}")
            return self.output_path
        
        except DocxTemplateError as e:
            print(f"错误: {str(e)}")
            raise
        except Exception as e:
            print(f"处理文档时发生未知错误: {str(e)}")
            raise DocxTemplateError(f"Unknown error while processing document: {str(e)}")

