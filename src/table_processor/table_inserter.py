"""
增强的表格插入器
支持：固定行数/动态行数策略、跳过列、填充数据
"""

from docx import Document
from typing import List, Dict, Optional, Any
import os

class EnhancedTableInserter:
    """增强版表格插入器"""
    
    def __init__(self, doc: Document):
        self.doc = doc
    
    def insert(self, 
               placeholder: str,
               table_template_path: str,
               raw_data: Optional[List[List[str]]] = None,
               transformations: Optional[List[Dict]] = None,
               metadata: Optional[Dict] = None,
               row_strategy: str = 'fixed_rows',
               skip_columns: Optional[List[int]] = None,
               header_rows: int = 1,
               location: str = 'body'):
        """
        插入表格（完整处理流程）
        
        Args:
            placeholder: 占位符名称
            table_template_path: 表格模板文件路径
            raw_data: 原始数据
            transformations: 数据转换规则
            metadata: 元数据
            row_strategy: 行策略 ('fixed_rows' 或 'dynamic_rows')
            skip_columns: 要跳过的列索引（模板中的固定列）
            header_rows: 表头行数
            location: 位置 ('body', 'header', 'footer')
        """
        if not os.path.exists(table_template_path):
            raise ValueError(f"Table template not found: {table_template_path}")
        
        from .data_transformer import TableDataTransformer
        
        transformer = TableDataTransformer()
        
        processed_data = raw_data
        if raw_data and transformations:
            processed_data = transformer.transform(raw_data, transformations, metadata)
        
        template_doc = Document(table_template_path)
        if not template_doc.tables:
            raise ValueError(f"No tables in template: {table_template_path}")
        
        template_table = template_doc.tables[0]
        
        if row_strategy == 'fixed_rows':
            self._fill_fixed_rows(template_table, processed_data, skip_columns, header_rows)
        elif row_strategy == 'dynamic_rows':
            self._fill_dynamic_rows(template_table, processed_data, skip_columns, header_rows)
        
        self._insert_to_document(placeholder, template_table, location)
    
    def _fill_fixed_rows(self, 
                        table: Any,
                        data: List[List[Any]],
                        skip_columns: Optional[List[int]],
                        header_rows: int):
        """
        固定行数模式：填充数据，跳过skip_columns和header_rows指定的列/行
        """
        if not data:
            return
        
        data_row_idx = 0
        for row_idx, row in enumerate(table.rows):
            if row_idx < header_rows:
                continue
            
            if data_row_idx >= len(data):
                break
            
            data_row = data[data_row_idx]
            data_col_idx = 0
            
            for col_idx, cell in enumerate(row.cells):
                if skip_columns and col_idx in skip_columns:
                    continue
                
                if data_col_idx < len(data_row):
                    value = data_row[data_col_idx]
                    self._set_cell_value(cell, str(value) if value else '')
                    data_col_idx += 1
            
            data_row_idx += 1
    
    def _fill_dynamic_rows(self,
                          table: Any,
                          data: List[List[Any]],
                          skip_columns: Optional[List[int]],
                          header_rows: int):
        """
        动态行数模式：删除所有数据行，根据数据量重新生成行
        """
        if not data:
            return
        
        while len(table.rows) > header_rows:
            table._tbl.remove(table.rows[-1]._tr)
        
        num_columns = len(table.rows[0].cells) if table.rows else 0
        
        for data_row in data:
            new_row = table.add_row()
            
            if len(new_row.cells) < num_columns:
                for _ in range(num_columns - len(new_row.cells)):
                    new_row.add_cell()
            
            data_col_idx = 0
            for col_idx, cell in enumerate(new_row.cells):
                if skip_columns and col_idx in skip_columns:
                    continue
                
                if data_col_idx < len(data_row):
                    value = data_row[data_col_idx]
                    self._set_cell_value(cell, str(value) if value else '')
                    data_col_idx += 1
    
    def _set_cell_value(self, cell: Any, value: str):
        """设置单元格值"""
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.text = value
                break
            if paragraph.runs:
                break
        else:
            cell.add_paragraph(value)
    
    def _insert_to_document(self, placeholder: str, table: Any, location: str):
        """将表格插入到文档占位符位置"""
        import sys
        sys.path.insert(0, str(os.path.dirname(__file__) + '/..'))
        from processor import PlaceholderFinder
        
        results = PlaceholderFinder.find_all_placeholders_in_location(
            self.doc, placeholder, location
        )
        
        if not results:
            raise ValueError(f"Placeholder not found: {placeholder}")
        
        for idx, paragraph in results:
            if paragraph._element.getparent() is None:
                continue
            PlaceholderFinder.replace_paragraph_with_element(
                paragraph, table._element, location
            )
